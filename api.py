"""
Vigilora AML — FastAPI Backend v2
Full feature parity with Python Streamlit dashboard.

Run:
    pip install fastapi uvicorn pandas openpyxl python-multipart python-docx
    python -m uvicorn api:app --reload --port 5000
"""

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
import pandas as pd
import numpy as np
import io
import warnings
import docx
from datetime import datetime
warnings.filterwarnings("ignore")

app = FastAPI(title="Vigilora AML API", version="2.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:5001",
        "http://127.0.0.1:5001",
        # ── Vercel production + previews ──────────────────────────────
        "https://vigilora-dashboard.vercel.app",
        "https://*.vercel.app",                    # covers all preview deployments
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── CONSTANTS ────────────────────────────────────────────────────────────────
HIGH_RISK_MCC = {
    '6012': 'Financial Institution Merchandise',
    '7995': 'Gambling / Betting',
    '6211': 'Securities Dealers',
    '5993': 'Cigar / Tobacco Stores',
    '5912': 'Drug Stores / Pharmacies',
    '7801': 'Government-Licensed Casinos',
    '6051': 'Non-Financial Institutions (Crypto/Currency)',
    '6010': 'Manual Cash Disbursements',
    '4829': 'Wire Transfers / Money Orders',
    '7800': 'Government-Licensed Gambling',
    '7802': 'State Lotteries',
}

FATF_HIGH_RISK = {
    'IRN': 'Iran',        'PRK': 'North Korea', 'MMR': 'Myanmar',
    'AFG': 'Afghanistan', 'IRQ': 'Iraq',        'SYR': 'Syria',
    'YEM': 'Yemen',       'LBY': 'Libya',       'SOM': 'Somalia',
    'VEN': 'Venezuela',   'NGA': 'Nigeria',     'PAK': 'Pakistan',
    'HTI': 'Haiti',       'RUS': 'Russia',      'BLR': 'Belarus',
    'CUB': 'Cuba',
}

# ─── HELPERS ──────────────────────────────────────────────────────────────────
def safe_col(df, *names):
    for n in names:
        if n in df.columns:
            return n
    return None

def fmt_amount(val):
    try:
        return f"${float(val):,.2f}"
    except (ValueError, TypeError):
        return str(val) if val is not None else 'N/A'

def normalize_mcc(val):
    try:
        return str(int(float(val)))
    except (ValueError, TypeError):
        return str(val).strip()

def load_excel(f) -> pd.DataFrame:
    try:
        df = pd.read_excel(f, engine='openpyxl')
    except Exception:
        df = pd.read_excel(f)
    df.columns = [str(c).strip() for c in df.columns]
    for col in df.select_dtypes(include=['object']).columns:
        df[col] = df[col].astype(str).replace({'nan': '', 'None': ''})
    return df

def safe_val(v):
    """Convert any pandas/numpy non-JSON-serializable value to a Python primitive."""
    if v is None:
        return None
    if isinstance(v, float) and (np.isnan(v) or np.isinf(v)):
        return None
    if isinstance(v, (np.integer,)):
        return int(v)
    if isinstance(v, (np.floating,)):
        return float(v)
    if isinstance(v, (pd.Timestamp, np.datetime64)):
        return str(v)
    if isinstance(v, (pd.NaT.__class__,)):
        return None
    # Catch-all: convert any remaining non-primitive to string
    if not isinstance(v, (bool, int, float, str, type(None))):
        return str(v)
    return v

def row_to_dict(row):
    return {k: safe_val(v) for k, v in row.items()}

# ─── SCORING ENGINE ───────────────────────────────────────────────────────────
def compute_risk_scores(df: pd.DataFrame) -> pd.DataFrame:
    idx = df.index
    n   = len(df)

    amt_col     = safe_col(df, 'Transaction Amount', 'Amount')
    mcc_col     = safe_col(df, 'Merchant Category Code', 'MCC')
    sender_c    = safe_col(df, 'SENDER ADDRESS COUNTRY', 'Sender Country')
    recv_c      = safe_col(df, 'RECEIVER ADDRESS COUNTRY', 'Receiver Country (Legacy)', 'Receiver Country', 'RECEIVER ADDRESS COUNTRY')
    card_col    = safe_col(df, 'CARD NUMBER', 'Card Number')
    recv_fn_col = safe_col(df, 'Receiver First Name', 'RECEIVER FIRST NAME')
    recv_ln_col = safe_col(df, 'Receiver Last Name', 'RECEIVER LAST NAME')

    total  = pd.Series(0, index=idx, dtype=int)
    flags  = [[] for _ in range(n)]
    bkdown = [[] for _ in range(n)]

    f_amt   = pd.Series(0.0, index=idx)
    f_mcc   = pd.Series(0, index=idx)
    f_cross = pd.Series(0, index=idx)
    f_fatf  = pd.Series(0, index=idx)
    f_round = pd.Series(0, index=idx)
    f_rvel  = pd.Series(0, index=idx)
    f_cvel  = pd.Series(0, index=idx)
    f_z     = pd.Series(0.0, index=idx)

    if amt_col:
        amt   = pd.to_numeric(df[amt_col], errors='coerce').fillna(0)
        f_amt = amt
        pts   = pd.Series(0, index=idx)
        pts[amt > 50000] = 25
        pts[(amt > 10000) & (amt <= 50000)] = 20
        pts[(amt > 5000)  & (amt <= 10000)] = 10
        total += pts
        for i, (p, a) in enumerate(zip(pts, amt)):
            if p == 25:   flags[i].append('Very Large Amount (>$50K)');  bkdown[i].append(f'Large Amount +{p}')
            elif p == 20: flags[i].append('Large Amount (>$10K)');       bkdown[i].append(f'Large Amount +{p}')
            elif p == 10: flags[i].append('Elevated Amount (>$5K)');     bkdown[i].append(f'Elevated Amount +{p}')

    if mcc_col:
        mcc_s  = df[mcc_col].apply(normalize_mcc)
        is_hr  = mcc_s.isin(HIGH_RISK_MCC.keys())
        f_mcc  = is_hr.astype(int)
        pts    = is_hr.astype(int) * 20
        total += pts
        for i, (flag, m, p) in enumerate(zip(is_hr, mcc_s, pts)):
            if flag:
                flags[i].append(f'High-Risk MCC {m} ({HIGH_RISK_MCC[m]})')
                bkdown[i].append(f'High-Risk MCC +{p}')
    else:
        mcc_s = pd.Series(['N/A'] * n, index=idx)

    if sender_c and recv_c:
        sc_vals  = df[sender_c].astype(str).str.upper().str.strip()
        rc_vals  = df[recv_c].astype(str).str.upper().str.strip()
        is_cross = sc_vals != rc_vals
        is_fs    = sc_vals.isin(FATF_HIGH_RISK.keys())
        is_fr    = rc_vals.isin(FATF_HIGH_RISK.keys())
        is_fatf  = is_fs | is_fr
        f_cross  = is_cross.astype(int)
        f_fatf   = is_fatf.astype(int)
        pts = (is_cross.astype(int)*10 + is_fatf.astype(int)*15 + (is_cross & is_fatf).astype(int)*5).clip(upper=30)
        total += pts
        sc_list = sc_vals.tolist()
        for i, (cross, fatf, p, r, s) in enumerate(zip(is_cross, is_fatf, pts, rc_vals, sc_list)):
            if fatf:
                cname = FATF_HIGH_RISK.get(str(r), '') or FATF_HIGH_RISK.get(str(s), str(r))
                flags[i].append(f'FATF High-Risk Country ({cname})')
                bkdown[i].append(f'FATF +{p}')
            elif cross:
                flags[i].append('Cross-Border Transaction')
                bkdown[i].append(f'Cross-Border +{p}')

    if amt_col:
        amt    = pd.to_numeric(df[amt_col], errors='coerce').fillna(0)
        is_1k  = (amt % 1000 == 0) & (amt >= 5000)
        is_5h  = (amt % 500  == 0) & (amt >= 2000) & ~is_1k
        f_round = is_1k.astype(int) * 2 + is_5h.astype(int)
        pts    = is_1k.astype(int) * 10 + is_5h.astype(int) * 5
        total += pts
        for i, (k, h, p) in enumerate(zip(is_1k, is_5h, pts)):
            if k:   flags[i].append('Round Amount ($1,000 multiple)'); bkdown[i].append(f'Round Amount +{p}')
            elif h: flags[i].append('Round Amount ($500 multiple)');   bkdown[i].append(f'Round Amount +{p}')

    if recv_fn_col and recv_ln_col:
        rkey  = df[recv_fn_col].astype(str) + '_' + df[recv_ln_col].astype(str)
        rfreq = rkey.map(rkey.value_counts())
        f_rvel = rfreq.fillna(0).astype(int)
        is_hrr = rfreq > 5
        pts    = is_hrr.astype(int) * 10
        total += pts
        for i, (flag, p, cnt) in enumerate(zip(is_hrr, pts, rfreq)):
            if flag:
                flags[i].append(f'High Receiver Velocity ({int(cnt)}x same receiver)')
                bkdown[i].append(f'Receiver Velocity +{p}')

    if card_col:
        cfreq  = df[card_col].map(df[card_col].value_counts())
        f_cvel = cfreq.fillna(0).astype(int)
        is_hrc = cfreq > 10
        pts    = is_hrc.astype(int) * 5
        total += pts
        for i, (flag, p, cnt) in enumerate(zip(is_hrc, pts, cfreq)):
            if flag:
                flags[i].append(f'High Card Usage ({int(cnt)}x same card)')
                bkdown[i].append(f'Card Velocity +{p}')

    if amt_col:
        amt  = pd.to_numeric(df[amt_col], errors='coerce').fillna(0)
        mean = amt.mean(); std = amt.std()
        if std > 0:
            z   = (amt - mean) / std
            f_z = z.round(2)
            anom = z.abs() > 2
            pts  = anom.astype(int) * 15
            total += pts
            for i, (flag, p, zv) in enumerate(zip(anom, pts, z)):
                if flag:
                    flags[i].append(f'Statistical Anomaly (Z={zv:.2f})')
                    bkdown[i].append(f'Z-Score Anomaly +{p}')

    total = total.clip(upper=100)
    level = total.apply(lambda s: 'HIGH' if s >= 61 else ('MEDIUM' if s >= 31 else 'LOW'))

    out = pd.DataFrame(index=idx)
    out['score']                 = total
    out['risk_level']            = level
    out['flags']                 = [' | '.join(f) if f else 'No flags' for f in flags]
    out['score_breakdown']       = [', '.join(b) if b else 'No risk points scored' for b in bkdown]
    out['feature_amount']        = f_amt
    out['feature_high_risk_mcc'] = f_mcc
    out['feature_cross_border']  = f_cross
    out['feature_fatf']          = f_fatf
    out['feature_round_amount']  = f_round
    out['feature_recv_velocity'] = f_rvel
    out['feature_card_velocity'] = f_cvel
    out['feature_zscore']        = f_z
    return out

def merge_results(df: pd.DataFrame, scores: pd.DataFrame) -> pd.DataFrame:
    r = df.copy()
    r['Risk Score']      = scores['score'].values
    r['Risk Level']      = scores['risk_level'].values
    r['Risk Flags']      = scores['flags'].values
    r['Score Breakdown'] = scores['score_breakdown'].values
    for c in [col for col in scores.columns if col.startswith('feature_')]:
        r[c] = scores[c].values
    r['Z-Score'] = scores['feature_zscore'].values
    return r

# ─── SAR GENERATION ───────────────────────────────────────────────────────────
def generate_sar_text(row: dict, total: int, avg_score: float) -> str:
    ref     = row.get('transaction_reference_number') or row.get('Transaction Reference') or 'N/A'
    amt_raw = None
    for k in row:
        if 'amount' in k.lower():
            amt_raw = row[k]; break
    amount  = fmt_amount(amt_raw)
    score   = row.get('Risk Score', 'N/A')
    level   = row.get('Risk Level', 'N/A')
    flags   = row.get('Risk Flags', 'N/A')
    bkdown  = row.get('Score Breakdown', 'N/A')
    z       = row.get('Z-Score', 0)
    date    = row.get('Transaction Date') or row.get('Date') or 'N/A'
    sender_country   = row.get('SENDER ADDRESS COUNTRY') or row.get('Sender Country') or 'N/A'
    receiver_country = row.get('RECEIVER ADDRESS COUNTRY') or row.get('Receiver Country') or row.get('Receiver Country (Legacy)') or 'N/A'
    mcc     = row.get('Merchant Category Code') or row.get('MCC') or 'N/A'
    merchant= row.get('Merchant Name') or 'N/A'
    card    = row.get('CARD NUMBER') or row.get('Card Number') or 'XXXX-XXXX-XXXX-XXXX'
    recv_fn = row.get('Receiver First Name') or row.get('RECEIVER FIRST NAME') or ''
    recv_ln = row.get('Receiver Last Name')  or row.get('RECEIVER LAST NAME')  or ''
    receiver_name = f"{recv_fn} {recv_ln}".strip() or 'N/A'

    now  = datetime.now()
    text = f"""
╔══════════════════════════════════════════════════════════════════════════════╗
║               VIGILORA AML — SUSPICIOUS ACTIVITY REPORT (SAR) DRAFT         ║
║                    ** NOT FOR FILING WITHOUT MLRO APPROVAL **                ║
╚══════════════════════════════════════════════════════════════════════════════╝

REPORT METADATA
───────────────────────────────────────────────────────────────────────────────
Report Generated     : {now.strftime('%Y-%m-%d %H:%M:%S')} UTC
SAR Reference        : SAR-{now.strftime('%Y%m%d')}-{str(ref)[:10]}
Status               : DRAFT — PENDING MLRO REVIEW
Reporting Institution: [FILL IN INSTITUTION NAME]
MLRO Name            : [FILL IN MLRO NAME]
Jurisdiction         : [FILL IN JURISDICTION]

SUBJECT TRANSACTION DETAILS
───────────────────────────────────────────────────────────────────────────────
Transaction Reference: {ref}
Transaction Date/Time: {date}
Transaction Amount   : {amount}
Merchant Name        : {merchant}
Merchant Category    : {mcc}
Card Number (masked) : {str(card)[:4]}XXXXXXXX{str(card)[-4:] if len(str(card)) >= 4 else '****'}
Sender Country       : {sender_country}
Receiver Country     : {receiver_country}
Receiver Name        : {receiver_name}

AI RISK ASSESSMENT
───────────────────────────────────────────────────────────────────────────────
AML Risk Score       : {score}/100
Risk Classification  : {level}
Z-Score Anomaly      : {float(z):.2f}σ  {'[ANOMALOUS — |Z|>2]' if abs(float(z)) > 2 else '[Within normal range]'}
Portfolio Context    : Batch avg score {avg_score:.1f}/100 across {total:,} transactions

Flags Triggered:
{chr(10).join(f'  → {f}' for f in str(flags).split(' | ')) if flags != 'No flags' else '  (None)'}

Score Breakdown:
{chr(10).join(f'  • {b}' for b in str(bkdown).split(', ')) if bkdown != 'No risk points scored' else '  (No points scored)'}

NARRATIVE
───────────────────────────────────────────────────────────────────────────────
On {date}, a transaction of {amount} (Reference: {ref}) was identified by the
Vigilora AML AI scoring engine as HIGH RISK with a score of {score}/100.

The transaction originated from {sender_country} and was directed to {receiver_country}.
{'This transaction involves a FATF high-risk jurisdiction, indicating elevated money laundering risk.' if receiver_country in FATF_HIGH_RISK else 'This is a cross-border transaction subject to enhanced due diligence.'}

{'The transaction amount exhibits a statistical anomaly (Z-Score: ' + str(float(z))[:5] + 'σ), significantly deviating from the portfolio average, indicating possible structuring or unusual activity.' if abs(float(z)) > 2 else ''}

The merchant category code {mcc} {'(' + HIGH_RISK_MCC.get(normalize_mcc(mcc), '') + ') is classified as high-risk under AML typologies.' if normalize_mcc(mcc) in HIGH_RISK_MCC else 'was reviewed as part of the overall risk assessment.'}

Based on the foregoing indicators, this transaction is recommended for enhanced
due diligence review and formal SAR filing consideration.

REQUIRED ACTIONS
───────────────────────────────────────────────────────────────────────────────
  [ ] MLRO to review and validate all risk indicators above
  [ ] Verify customer identity and source of funds
  [ ] Check sanctions and PEP screening results
  [ ] Determine whether formal SAR filing is required
  [ ] If filing: Submit to regulatory authority within statutory deadline
  [ ] Document decision and rationale in case management system

DISCLAIMER
───────────────────────────────────────────────────────────────────────────────
This SAR draft is auto-generated by the Vigilora AI AML system and must be
reviewed, validated, and approved by an authorized Money Laundering Reporting
Officer (MLRO) before filing. Do not share without MLRO authorization.

═════════════════════════ END OF SAR DRAFT ═════════════════════════
""".strip()
    return text

# ─── IN-MEMORY STORE (per request — stateless) ───────────────────────────────
# We re-score on each /analyze call. For SAR/export we store per session via
# a simple request-scoped cache keyed by upload hash.
_result_cache: dict[str, pd.DataFrame] = {}

# ─── ENDPOINTS ───────────────────────────────────────────────────────────────
@app.get("/")
def root():
    return {"status": "Vigilora AML API v3.4", "docs": "/docs"}

@app.get("/api/health")
def health():
    return {"status": "ok"}


@app.post("/api/analyze")
async def analyze(file: UploadFile = File(...)):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(400, "Only .xlsx / .xls files are supported.")

    contents = await file.read()
    try:
        df_raw    = load_excel(io.BytesIO(contents))
        # Rule scores
        scores    = compute_risk_scores(df_raw)
        
        # Inject ML Hybrid scoring if models are trained
        try:
            from app.ml.predict import models_ready, merge_ml_with_rules, score_features
            from app.ml.features import build_features
            if models_ready():
                X, _ = build_features(df_raw)
                ml_sc = score_features(X)
                blended = merge_ml_with_rules(scores['score'], ml_sc)
                scores['score'] = blended
                scores['risk_level'] = blended.apply(lambda s: 'HIGH' if s >= 70 else ('MEDIUM' if s >= 40 else 'LOW'))
        except Exception as e:
            print(f"ML injection failed, falling back to rules: {e}")
            
        df_result = merge_results(df_raw, scores)
    except Exception as e:
        raise HTTPException(500, f"Scoring failed: {e}")

    # Cache for SAR/export endpoints
    cache_key = str(hash(contents))
    _result_cache[cache_key] = df_result

    total      = len(df_result)
    high_n     = int((df_result['Risk Level'] == 'HIGH').sum())
    med_n      = int((df_result['Risk Level'] == 'MEDIUM').sum())
    low_n      = int((df_result['Risk Level'] == 'LOW').sum())
    avg_score  = float(df_result['Risk Score'].mean())

    sc_col  = safe_col(df_result, 'SENDER ADDRESS COUNTRY', 'Sender Country')
    rc_col  = safe_col(df_result, 'RECEIVER ADDRESS COUNTRY', 'Receiver Country (Legacy)', 'Receiver Country')
    amt_col = safe_col(df_result, 'Transaction Amount', 'Amount')
    ref_col = safe_col(df_result, 'transaction_reference_number', 'Transaction Reference')
    date_col= safe_col(df_result, 'Transaction Date', 'Date')
    mcc_col = safe_col(df_result, 'Merchant Category Code', 'MCC')

    cross_count = 0
    cross_pct   = 0.0
    if sc_col and rc_col and total > 0:
        cross_count = int((df_result[sc_col].astype(str).str.upper() != df_result[rc_col].astype(str).str.upper()).sum())
        cross_pct   = round(cross_count / total * 100, 1)

    fatf_pct = 0.0
    if rc_col and total > 0:
        fatf_pct = round(df_result[rc_col].astype(str).str.upper().isin(FATF_HIGH_RISK.keys()).sum() / total * 100, 1)

    kpis = {
        "total": total, "high": high_n, "medium": med_n, "low": low_n,
        "avgScore": round(avg_score, 1), "sarCandidates": high_n,
        "customerRiskPct":    round(high_n / total * 100, 1) if total > 0 else 0,
        "transactionRiskPct": round((high_n + med_n) / total * 100, 1) if total > 0 else 0,
        "geographicRiskPct":  cross_pct,
        "fatfPct": fatf_pct,
    }

    # --- Multi-granularity Trend ---
    trend_daily = []
    trend_weekly = []
    trend_monthly = []
    default_view = "daily"

    if date_col:
        try:
            df_result['_date'] = pd.to_datetime(df_result[date_col], errors='coerce')
            df_dated = df_result.dropna(subset=['_date']).copy()

            if not df_dated.empty:
                date_range_days = (df_dated['_date'].max() - df_dated['_date'].min()).days

                # Determine default view based on data span
                if date_range_days >= 60:
                    default_view = "monthly"
                elif date_range_days >= 14:
                    default_view = "weekly"
                else:
                    default_view = "daily"

                agg_cols = {'Risk Score': 'mean'}
                high_agg = {'Risk Level': lambda x: (x == 'HIGH').sum()}
                if amt_col:
                    agg_cols[amt_col] = 'mean'

                # Daily
                grp_d = df_dated.groupby(df_dated['_date'].dt.date)
                sc_d  = grp_d['Risk Score'].mean()
                hi_d  = grp_d['Risk Level'].apply(lambda x: (x == 'HIGH').sum())
                am_d  = grp_d[amt_col].mean() if amt_col else None
                for day, score in sc_d.items():
                    p = {"label": str(day), "score": round(float(score), 1),
                         "highCount": int(hi_d.get(day, 0))}
                    if am_d is not None:
                        p["avgAmount"] = round(float(am_d.get(day, 0)), 2)
                    trend_daily.append(p)

                # Weekly
                df_dated['_week'] = df_dated['_date'].dt.to_period('W').dt.start_time.dt.date
                grp_w = df_dated.groupby('_week')
                sc_w  = grp_w['Risk Score'].mean()
                hi_w  = grp_w['Risk Level'].apply(lambda x: (x == 'HIGH').sum())
                am_w  = grp_w[amt_col].mean() if amt_col else None
                for wk, score in sc_w.items():
                    p = {"label": f"W/{str(wk)[5:]}", "score": round(float(score), 1),
                         "highCount": int(hi_w.get(wk, 0))}
                    if am_w is not None:
                        p["avgAmount"] = round(float(am_w.get(wk, 0)), 2)
                    trend_weekly.append(p)

                # Monthly
                df_dated['_month'] = df_dated['_date'].dt.to_period('M').dt.start_time.dt.date
                grp_m = df_dated.groupby('_month')
                sc_m  = grp_m['Risk Score'].mean()
                hi_m  = grp_m['Risk Level'].apply(lambda x: (x == 'HIGH').sum())
                am_m  = grp_m[amt_col].mean() if amt_col else None
                for mo, score in sc_m.items():
                    p = {"label": str(mo)[:7], "score": round(float(score), 1),
                         "highCount": int(hi_m.get(mo, 0))}
                    if am_m is not None:
                        p["avgAmount"] = round(float(am_m.get(mo, 0)), 2)
                    trend_monthly.append(p)

        except Exception:
            pass

    # Fallback batches if no date column
    if not trend_daily:
        bsz = max(1, total // 8)
        for i in range(min(8, total)):
            chunk = df_result.iloc[i*bsz:(i+1)*bsz]
            p = {"label": f"Batch {i+1}", "score": round(float(chunk['Risk Score'].mean()), 1),
                 "highCount": int((chunk['Risk Level'] == 'HIGH').sum())}
            if amt_col:
                p["avgAmount"] = round(float(pd.to_numeric(chunk[amt_col], errors='coerce').mean()), 2)
            trend_daily.append(p)
        trend_weekly  = trend_daily
        trend_monthly = trend_daily

    # Keep backward-compat `trend` key pointing to daily
    trend = [{"day": p["label"], "score": p["score"]} for p in trend_daily]

    # Drop temp column before building any allFields dicts
    if '_date' in df_result.columns:
        df_result = df_result.drop(columns=['_date'])

    # --- Transactions (top 200) ---
    tx_cols = [c for c in [ref_col, date_col, amt_col, sc_col, rc_col,
               'Risk Score', 'Risk Level', 'Risk Flags', 'Score Breakdown', 'Z-Score',
               safe_col(df_result, 'Merchant Name'), mcc_col]
               if c and c in df_result.columns]
    transactions = [row_to_dict(row) for _, row in df_result[tx_cols].head(200).iterrows()]

    # --- Overview: risk distribution buckets ---
    score_buckets = {}
    for v in df_result['Risk Score']:
        b = f"{(int(v)//10)*10}-{(int(v)//10)*10+10}"
        score_buckets[b] = score_buckets.get(b, 0) + 1
    histogram = [{"range": k, "count": v} for k, v in sorted(score_buckets.items())]

    risk_summary = []
    for lvl in ['HIGH', 'MEDIUM', 'LOW']:
        sub = df_result[df_result['Risk Level'] == lvl]['Risk Score']
        if not sub.empty:
            risk_summary.append({
                "level": lvl, "count": int(len(sub)),
                "avg": round(float(sub.mean()), 1),
                "max": int(sub.max()), "min": int(sub.min()),
            })

    # --- High & Medium alert cards ---
    def make_alert_cards(rows, max_n=40):
        cards = []
        for _, row in rows.head(max_n).iterrows():
            ref = row[ref_col] if ref_col else str(row.name)
            amt = fmt_amount(row[amt_col]) if amt_col else 'N/A'
            flags = row['Risk Flags']
            flag_count = flags.count('|') + 1 if flags != 'No flags' else 0
            cards.append({
                "ref": str(ref), "amount": amt,
                "score": int(row['Risk Score']), "level": str(row['Risk Level']),
                "flags": str(flags), "breakdown": str(row['Score Breakdown']),
                "zScore": round(float(row.get('Z-Score', 0)), 2),
                "flagCount": flag_count,
            })
        return cards

    high_alerts   = make_alert_cards(df_result[df_result['Risk Level'] == 'HIGH'])
    medium_alerts = make_alert_cards(df_result[df_result['Risk Level'] == 'MEDIUM'])

    # --- Anomalies ---
    anomalies = []
    anom_df = df_result[df_result['Z-Score'].abs() > 2].sort_values('Z-Score', key=abs, ascending=False).head(20)
    for _, row in anom_df.iterrows():
        ref = row[ref_col] if ref_col else str(row.name)
        anomalies.append({
            "ref": str(ref),
            "zScore": round(float(row['Z-Score']), 2),
            "score": int(row['Risk Score']),
            "flags": str(row['Risk Flags'])[:80],
        })

    # --- Merchant data ---
    merchant_data = []
    if mcc_col:
        mcc_norm = df_result[mcc_col].apply(normalize_mcc)
        top_mcc  = mcc_norm.value_counts().head(12).reset_index()
        top_mcc.columns = ['mcc', 'count']
        for _, row in top_mcc.iterrows():
            sub = df_result[mcc_norm == row['mcc']]
            merchant_data.append({
                "mcc":         str(row['mcc']),
                "description": HIGH_RISK_MCC.get(str(row['mcc']), 'Standard Merchant'),
                "count":       int(row['count']),
                "isHighRisk":  str(row['mcc']) in HIGH_RISK_MCC,
                "avgScore":    round(float(sub['Risk Score'].mean()), 1),
            })

    # --- FATF transactions ---
    fatf_transactions = []
    if rc_col:
        fatf_df = df_result[df_result[rc_col].astype(str).str.upper().isin(FATF_HIGH_RISK.keys())].head(30)
        for _, row in fatf_df.iterrows():
            ctry = str(row[rc_col]).upper()
            fatf_transactions.append({
                "country":     ctry,
                "countryName": FATF_HIGH_RISK.get(ctry, ctry),
                "amount":      fmt_amount(row[amt_col]) if amt_col else 'N/A',
                "score":       int(row['Risk Score']),
                "level":       str(row['Risk Level']),
                "ref":         str(row[ref_col]) if ref_col else 'N/A',
            })

    # --- Country bar data ---
    country_data = []
    if rc_col:
        top_c = df_result[rc_col].value_counts().head(15).reset_index()
        top_c.columns = ['country', 'count']
        for _, row in top_c.iterrows():
            country_data.append({
                "country": str(row['country']),
                "count": int(row['count']),
                "isFatf": str(row['country']).upper() in FATF_HIGH_RISK,
            })

    # --- Network Graph Data ---
    network = {"nodes": [], "edges": []}
    if ref_col:
        # We need sender and receiver. If not explicitly present, use ref as a proxy, or mock it if missing.
        # Let's see if we have source/dest accounts. Common names: 'Sender Account', 'Receiver Account', 'Originator Account', 'Beneficiary Account'
        sn_col = safe_col(df_result, 'Sender Account') or safe_col(df_result, 'Originator Account')
        rn_col = safe_col(df_result, 'Receiver Account') or safe_col(df_result, 'Beneficiary Account')

        if sn_col and rn_col:
            # Get edges (transactions)
            edges_df = df_result.groupby([sn_col, rn_col])['Risk Score'].max().reset_index()
            # Sort by risk score descending, take top 25 edges to avoid clutter
            edges_df = edges_df.sort_values('Risk Score', ascending=False).head(25)

            node_set = set(edges_df[sn_col]).union(set(edges_df[rn_col]))
            
            # Build node dictionaries
            node_map = {}
            for n in node_set:
                # Find max risk for this node across all its transactions
                sn_max = df_result[df_result[sn_col] == n]['Risk Score'].max() if n in df_result[sn_col].values else 0
                rn_max = df_result[df_result[rn_col] == n]['Risk Score'].max() if n in df_result[rn_col].values else 0
                max_score = max(sn_max, rn_max)
                risk_level = 'high' if max_score >= 61 else 'medium' if max_score >= 31 else 'low'
                node_map[n] = {"id": str(n), "label": str(n)[:8], "risk": risk_level}
            
            network["nodes"] = list(node_map.values())
            network["edges"] = [{"from": str(r[sn_col]), "to": str(r[rn_col])} for _, r in edges_df.iterrows()]
        else:
            # Fallback if no explicit sender/receiver account columns: just group by Merchant or Country 
            # to show a different kind of relationship, or return empty to let frontend use mock.
            # Returning empty nodes array signals frontend to handle it.
            pass

    # --- SAR candidates ---
    sar_candidates = []
    high_sar = df_result[df_result['Risk Level'] == 'HIGH'].copy()
    for i, (idx_val, row) in enumerate(high_sar.iterrows()):
        ref = row[ref_col] if ref_col else f"Row {idx_val}"
        flags_short = str(row['Risk Flags'])[:60] + ('...' if len(str(row['Risk Flags'])) > 60 else '')
        sar_candidates.append({
            "index": i,
            "rowIndex": int(idx_val),
            "ref": str(ref),
            "score": int(row['Risk Score']),
            "zScore": round(float(row.get('Z-Score', 0)), 2),
            "flagCount": str(row['Risk Flags']).count('|') + 1 if row['Risk Flags'] != 'No flags' else 0,
            "flagsShort": flags_short,
            "breakdown": str(row['Score Breakdown']),
            "allFields": row_to_dict(row),
        })

    # --- Alerts ---
    alerts = []
    if fatf_transactions:
        alerts.append({"id": "fatf", "icon": "globe", "title": f"{len(fatf_transactions)} FATF high-risk jurisdiction transactions detected", "severity": "danger", "timestamp": "Just now"})
    if high_n > 0:
        alerts.append({"id": "sar", "icon": "clipboard", "title": f"{high_n} SAR candidate(s) pending MLRO review", "severity": "warning", "timestamp": "Just now"})
    if anomalies:
        alerts.append({"id": "anomaly", "icon": "bar-chart", "title": f"{len(anomalies)} statistical anomalies detected (|Z|>2)", "severity": "warning", "timestamp": "Just now"})
    alerts.append({"id": "kyc", "icon": "alert", "title": "KYC Verification required for HIGH risk account holders", "severity": "danger", "timestamp": "Just now"})

    return JSONResponse({
        "cacheKey":        cache_key,
        "kpis":            kpis,
        "trend":           trend,
        "transactions":    transactions,
        "alerts":          alerts,
        "countryData":     country_data,
        "columns":         tx_cols,
        "histogram":       histogram,
        "riskSummary":     risk_summary,
        "highAlerts":      high_alerts,
        "mediumAlerts":    medium_alerts,
        "anomalies":       anomalies,
        "merchantData":    merchant_data,
        "fatfTransactions": fatf_transactions,
        "sarCandidates":   sar_candidates,
        "crossBorderCount": cross_count,
        "trendDaily":      trend_daily,
        "trendWeekly":     trend_weekly,
        "trendMonthly":    trend_monthly,
        "defaultView":     default_view,
        "network":         network,
    })


@app.get("/api/sar/{cache_key}/{row_index}")
async def get_sar(cache_key: str, row_index: int):
    """Generate full SAR narrative for a specific HIGH risk transaction (JSON formatting)."""
    df = _result_cache.get(cache_key)
    if df is None:
        raise HTTPException(404, "Session expired. Please re-upload the file.")
    high_sar = df[df['Risk Level'] == 'HIGH'].copy()
    if row_index >= len(high_sar):
        raise HTTPException(404, "SAR candidate index out of range.")
    row = high_sar.iloc[row_index]
    total     = len(df)
    avg_score = float(df['Risk Score'].mean())
    text = generate_sar_text(row.to_dict(), total, avg_score)
    return JSONResponse({"text": text})


@app.get("/api/sar/download/{cache_key}/{row_index}")
async def download_sar_docx(cache_key: str, row_index: int):
    """Generate and download SAR narrative as a standard Word (.docx) file."""
    df = _result_cache.get(cache_key)
    if df is None:
        raise HTTPException(404, "Session expired.")
    high_sar = df[df['Risk Level'] == 'HIGH'].copy()
    if row_index >= len(high_sar):
        raise HTTPException(404, "Index out of range.")
    
    row = high_sar.iloc[row_index]
    text = generate_sar_text(row.to_dict(), len(df), float(df['Risk Score'].mean()))
    
    # Create Word Document
    doc = docx.Document()
    
    # Add Header
    heading = doc.add_heading('SUSPICIOUS ACTIVITY REPORT (SAR)', 0)
    heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("_" * 60)
    
    # Add Content body
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('#'):
            # Convert markdown headings to Docx headings
            level = len(line) - len(line.lstrip('#'))
            doc.add_heading(line.lstrip('#').strip(), level=min(level, 4))
            continue
            
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            line = line[2:]
        else:
            p = doc.add_paragraph()
            
        # Parse **bold** markdown dynamically
        parts = line.split('**')
        for i, part in enumerate(parts):
            if not part: continue
            run = p.add_run(part)
            if i % 2 != 0: # inside bold markers
                run.bold = True
                
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    ref_val = str(row.get('Reference', row.get('Transaction Ref', f"TX_{row_index}")))
    fname = f"SAR_{ref_val}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )


@app.get("/api/export/{cache_key}")
async def export_excel(cache_key: str):
    """Export full analysis as Excel with 4 sheets."""
    df = _result_cache.get(cache_key)
    if df is None:
        raise HTTPException(404, "Session expired. Please re-upload the file.")

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='All Transactions', index=False)
        df[df['Risk Level'] == 'HIGH'].to_excel(writer, sheet_name='High Risk', index=False)
        df[df['Risk Level'] == 'MEDIUM'].to_excel(writer, sheet_name='Medium Risk', index=False)
        ml_cols = [c for c in df.columns if c.startswith('feature_')]
        df[ml_cols + ['Risk Score', 'Risk Level']].to_excel(writer, sheet_name='ML Training Data', index=False)
    buf.seek(0)

    fname = f"AML_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )
