from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
import io
import os
import docx
from datetime import datetime
import pandas as pd
from app.controllers.scoring_controller import load_excel, compute_risk_scores, merge_results
from app.controllers.sar_controller import generate_sar_text
from app.controllers.analysis_controller import build_analysis_response
from app.models.constants import _result_cache

router = APIRouter(tags=["Legacy Dashboard"])

@router.get("/")
def root():
    return {"status": "Vigilora AML API v3.9", "docs": "/docs"}

@router.get("/api/health")
@router.get("/health")
def health():
    return {"status": "ok v10"}

@router.get("/api/analyze/stored")
@router.get("/analyze/stored")
async def analyze_stored():
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    file_path = os.path.join(base_dir, "aml_dashboard_export_2026-02-16.xlsx")
    if not os.path.exists(file_path):
        raise HTTPException(404, "Stored Excel file not found on the server.")
    
    from fastapi.concurrency import run_in_threadpool
    def read_file():
        with open(file_path, "rb") as f:
            return f.read()
            
    contents = await run_in_threadpool(read_file)
        
    class MockUploadFile:
        def __init__(self, content_bytes):
            self.filename = "aml_dashboard_export_2026-02-16.xlsx"
            self.content_bytes = content_bytes
        async def read(self):
            return self.content_bytes
            
    return await analyze(file=MockUploadFile(contents))

@router.post("/api/analyze")
@router.post("/analyze")
async def analyze(file: UploadFile = File(...)):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(400, "Only .xlsx / .xls files are supported.")

    contents = await file.read()
    
    def process_analysis(file_contents):
        df_raw    = load_excel(io.BytesIO(file_contents))
        # Rule scores
        scores    = compute_risk_scores(df_raw)
        
        # Inject ML Hybrid scoring if models are trained
        try:
            from app.ml.predict import models_ready, merge_ml_with_rules, score_features
            from app.ml.features import build_features
            from app.ml.explain import explain_batch

            if models_ready():
                X, _ = build_features(df_raw)
                ml_sc = score_features(X)
                blended = merge_ml_with_rules(scores['score'], ml_sc)
                
                # --- ENTERPRISE SHAP ALERTS ---
                shap_explanations = explain_batch(X, top_n=3)
                SHAP_IMPACT_MIN_THRESHOLD = 1.5
                
                for i, exps in enumerate(shap_explanations):
                    if not exps: continue
                    
                    idx = blended.index[i]
                    shap_penalty = 0
                    alert_reasons = []
                    
                    for e in exps:
                        if e['impact'] >= SHAP_IMPACT_MIN_THRESHOLD:
                            penalty = min(int((e['impact'] - 1.0) * 15), 30) 
                            if penalty > 0:
                                shap_penalty += penalty
                                alert_reasons.append(f"AI Risk Indicator: {e['label']} (+{penalty} pts)")
                    
                    if shap_penalty > 0:
                        shap_penalty = min(shap_penalty, 40)
                        blended.iloc[i] = min(blended.iloc[i] + shap_penalty, 100)
                        
                        curr_flags = scores.at[idx, 'flags']
                        new_flags = " | ".join(alert_reasons)
                        if curr_flags == 'No flags':
                            scores.at[idx, 'flags'] = new_flags
                        else:
                            scores.at[idx, 'flags'] = str(curr_flags) + " | " + new_flags
                            
                        curr_bk = scores.at[idx, 'score_breakdown']
                        new_bk = f"ML Interpretability Penalty +{shap_penalty}"
                        if curr_bk == 'No risk points scored':
                            scores.at[idx, 'score_breakdown'] = new_bk
                        elif new_bk not in str(curr_bk):
                            scores.at[idx, 'score_breakdown'] = str(curr_bk) + ", " + new_bk

                scores['score'] = blended
                scores['risk_level'] = blended.apply(lambda s: 'HIGH' if s >= 70 else ('MEDIUM' if s >= 40 else 'LOW'))
        except Exception as e:
            print(f"ML injection failed, falling back to rules: {e}")
            
        df_result = merge_results(df_raw, scores)
        
        # Cache for SAR/export endpoints
        cache_key = str(hash(file_contents))
        _result_cache[cache_key] = df_result

        return build_analysis_response(df_result, cache_key)

    try:
        from fastapi.concurrency import run_in_threadpool
        analysis_response = await run_in_threadpool(process_analysis, contents)
    except Exception as e:
        raise HTTPException(500, f"Scoring failed: {e}")

    return JSONResponse(analysis_response)


@router.get("/api/sar/{cache_key}/{row_index}")
@router.get("/sar/{cache_key}/{row_index}")
async def get_sar(cache_key: str, row_index: int):
    """Generate full SAR narrative for a specific HIGH risk transaction (JSON formatting)."""
    df = _result_cache.get(cache_key)
    if df is None:
        raise HTTPException(404, "Session expired. Please re-upload the file.")
    high_sar = df[df['Risk Level'] == 'HIGH'].copy()
    if row_index >= len(high_sar):
        raise HTTPException(404, "SAR candidate index out of range.")
    row = high_sar.iloc[row_index]
    total     = len(df)
    avg_score = float(df['Risk Score'].mean())
    text = generate_sar_text(row.to_dict(), total, avg_score)
    return JSONResponse({"text": text})


@router.get("/api/sar/download/{cache_key}/{row_index}")
@router.get("/sar/download/{cache_key}/{row_index}")
async def download_sar_docx(cache_key: str, row_index: int):
    """Generate and download SAR narrative as a standard Word (.docx) file."""
    df = _result_cache.get(cache_key)
    if df is None:
        raise HTTPException(404, "Session expired.")
    high_sar = df[df['Risk Level'] == 'HIGH'].copy()
    if row_index >= len(high_sar):
        raise HTTPException(404, "Index out of range.")
    
    row = high_sar.iloc[row_index]
    text = generate_sar_text(row.to_dict(), len(df), float(df['Risk Score'].mean()))
    
    doc = docx.Document()
    heading = doc.add_heading('SUSPICIOUS ACTIVITY REPORT (SAR)', 0)
    heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("_" * 60)
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            doc.add_heading(line.lstrip('#').strip(), level=min(level, 4))
            continue
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            line = line[2:]
        else:
            p = doc.add_paragraph()
            
        parts = line.split('**')
        for i, part in enumerate(parts):
            if not part: continue
            run = p.add_run(part)
            if i % 2 != 0:
                run.bold = True
                
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    ref_val = str(row.get('Reference', row.get('Transaction Ref', f"TX_{row_index}")))
    fname = f"SAR_{ref_val}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )


@router.get("/api/export/{cache_key}")
@router.get("/export/{cache_key}")
async def export_excel(cache_key: str):
    """Export full analysis as Excel with 4 sheets."""
    df = _result_cache.get(cache_key)
    if df is None:
        raise HTTPException(404, "Session expired. Please re-upload the file.")

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='All Transactions', index=False)
        df[df['Risk Level'] == 'HIGH'].to_excel(writer, sheet_name='High Risk', index=False)
        df[df['Risk Level'] == 'MEDIUM'].to_excel(writer, sheet_name='Medium Risk', index=False)
        ml_cols = [c for c in df.columns if c.startswith('feature_')]
        df[ml_cols + ['Risk Score', 'Risk Level']].to_excel(writer, sheet_name='ML Training Data', index=False)
    buf.seek(0)

    fname = f"AML_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )
